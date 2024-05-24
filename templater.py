from docx import Document
from docxtpl import DocxTemplate
import re
import shutil
import os
import argparse

parser = argparse.ArgumentParser(description='Get attachments of the file')
parser.add_argument('--file', '-f', type=str, required=True, help='File')
parser.add_argument('--output', '-o', type=str, required=True, help='Output file')

# Parse command line arguments
args = parser.parse_args()

def clear_section_content(document_path, start_delimiter, end_delimiter):

    # Make a copy of the original document if the modified file doesn't exist yet
    if not os.path.exists(document_path):
        shutil.copy2(document_path, document_path)

    # Load the modified document
    doc = Document(document_path)

    # Clear the content of the specified section
    is_section = False
    paragraphs_to_delete = []

    for i, paragraph in enumerate(doc.paragraphs):
        if start_delimiter in paragraph.text:
            is_section = True
        elif end_delimiter in paragraph.text:
            is_section = False

        if is_section:
            paragraphs_to_delete.append(i)
            paragraph._element.clear_content()  # Clear content including numbering

    # Delete the paragraphs in reverse order to avoid shifting indices
    for index in reversed(paragraphs_to_delete):
        p = doc.paragraphs[index]
        p._element.getparent().remove(p._element)

    # Save the modified document
    doc.save(document_path)

    return document_path

clear_section_content(args.file, "{{ BEGIN_FIRST_READING }}", "{{ END_FIRST_READING }}")
clear_section_content(args.file, "{{ BEGIN_SECOND_READING }}", "{{ END_SECOND_READING }}")
clear_section_content(args.file, "{{ BEGIN_THIRD_READING }}", "{{ END_THIRD_READING }}")
clear_section_content(args.file, "{{ BEGIN_UNASSIGNED }}", "{{ END_UNASSIGNED }}")
clear_section_content(args.file, "{{ BEGIN_ANNOUNCEMENT }}", "{{ END_ANNOUNCEMENT }}")

print(f"Modified document saved")

template = DocxTemplate(args.file)

context = {
    "END_FIRST_READING" : "{first_reading_content}",
    "END_SECOND_READING" : "{second_reading_content}",
    "END_THIRD_READING" : "{third_reading_content}",
    "END_UNASSIGNED" : "${UnassignedContent}",
    "END_ANNOUNCEMENT" : "{announcement_content}",
}

template.render(context)

template.save(args.output)

print(f"Rendered successfully!")
